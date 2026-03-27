import spacy
import PyPDF2
import docx

# ...

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])
from django.conf import settings
import os

# Load NLP model
try:
    nlp = spacy.load("en_core_web_sm")
except:
    nlp = None # Handle case where model isn't downloaded yet

def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_skills(text, job_role_skills):
    """
    Simple skill extraction based on keyword matching.
    """
    text_lower = text.lower()
    found_skills = []
    missing_skills = []
    
    # Process text with spacy for lemmatization if needed, 
    # but for known technical terms, direct matching is often more robust for basics.
    # doc = nlp(text_lower) 
    
    for skill in job_role_skills:
        # Check if the skill exists in the text (case insensitive)
        # Using simple substring check for now, can be improved with regex
        if skill.lower() in text_lower:
            found_skills.append(skill)
        else:
            missing_skills.append(skill)
            
    return found_skills, missing_skills
